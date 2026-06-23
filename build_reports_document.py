from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Documentacion_general_informes_y_contraste_salesforce.md"
OUTPUT = ROOT / "Documentacion_general_informes_y_contraste_salesforce.docx"


def set_cell_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles.add_style("CodexTitle", WD_STYLE_TYPE.PARAGRAPH)
    title.base_style = normal
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)

    subtitle = styles.add_style("CodexSubtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.base_style = normal
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(10)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    subtitle.paragraph_format.space_after = Pt(6)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.15

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.15

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.line_spacing = 1.15

    code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code.base_style = normal
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.line_spacing = 1.0


def add_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("HR Motor · Informes · 2026-06-23")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def is_major_section(text: str) -> bool:
    return bool(re.match(r"^\d+\.\s", text))


def is_minor_section(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\.\s", text))


def is_sub_minor_section(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\.\d+\.\s", text))


def emit_heading(document: Document, text: str) -> None:
    if is_sub_minor_section(text):
        document.add_paragraph(text, style="Heading 3")
        return

    if is_minor_section(text):
        document.add_paragraph(text, style="Heading 2")
        return

    if is_major_section(text):
        if any(p.text.strip() for p in document.paragraphs):
            paragraph = document.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
        document.add_paragraph(text, style="Heading 1")
        return

    document.add_paragraph(text, style="Heading 2")


def build_document() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_page(document)
    configure_styles(document)
    add_footer(document)

    in_code = False

    for line in lines:
        stripped = line.rstrip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            paragraph = document.add_paragraph(stripped, style="CodeBlock")
            set_cell_shading(paragraph, "F4F6F9")
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            document.add_paragraph(stripped[2:].strip(), style="CodexTitle")
            continue

        if stripped.startswith("Versión:") or stripped.startswith("Proyecto:"):
            document.add_paragraph(stripped, style="CodexSubtitle")
            continue

        if stripped.startswith("## "):
            emit_heading(document, stripped[3:].strip())
            continue

        if stripped.startswith("### "):
            document.add_paragraph(stripped[4:].strip(), style="Heading 3")
            continue

        if re.match(r"^\d+\.\s", stripped):
            document.add_paragraph(stripped, style="List Number")
            continue

        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        paragraph = document.add_paragraph(style="Normal")
        run = paragraph.add_run(stripped)
        if stripped.endswith(":"):
            run.bold = True

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()

